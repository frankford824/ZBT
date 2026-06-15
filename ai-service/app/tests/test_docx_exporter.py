from __future__ import annotations

import base64
import json
from io import BytesIO
from zipfile import ZipFile

import fitz
import pytest
from docx import Document

from app.pipelines.export.docx_exporter import (
    _attachment_content,
    _replace_template_fields,
    _safe_zip_path,
    _validate_pdf_output,
    export_bid_docx,
    export_bid_pdf,
    export_bid_zip,
)
from app.schemas.export import ExportAttachment, ExportChapter, ExportLayoutOptions, ExportPart


def test_export_bid_docx_applies_master_layout(tmp_path, monkeypatch) -> None:
    monkeypatch.delenv("BID_EXPORT_TEMPLATE_PATH", raising=False)
    output = tmp_path / "bid.docx"
    layout = ExportLayoutOptions(watermark_text="内部评审", generated_at="2026-06-11")
    chapters = [
        ExportChapter(
            title="项目实施方案",
            plain_text=(
                "# 总体安排\n"
                "| 项目 | 内容 |\n"
                "| --- | --- |\n"
                "| 工期 | 90天 |\n\n"
                "- 进度控制\n"
                "1. 质量控制"
            ),
        )
    ]

    export_bid_docx("智慧交通平台", "技术标", chapters, output, layout=layout)

    document = Document(output)
    paragraph_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    header_text = "\n".join(paragraph.text for paragraph in document.sections[0].header.paragraphs)
    footer_text = "\n".join(paragraph.text for paragraph in document.sections[0].footer.paragraphs)
    package_xml = _package_xml(output)

    assert "智慧交通平台" in paragraph_text
    assert "投标文件" in paragraph_text
    assert "目录" in paragraph_text
    assert "技术标" in paragraph_text
    assert "项目实施方案" in paragraph_text
    assert "总体安排" in paragraph_text
    assert "智慧交通平台 - 技术标" in header_text
    assert "智标通投标文件导出" in footer_text
    assert len(document.tables) == 1
    assert document.tables[0].cell(1, 0).text == "工期"
    assert "TOC" in package_xml
    assert "updateFields" in package_xml
    assert "PAGE" in package_xml
    assert "NUMPAGES" in package_xml
    assert "内部评审" in package_xml


def test_export_bid_zip_uses_master_layout_for_each_part(tmp_path, monkeypatch) -> None:
    monkeypatch.delenv("BID_EXPORT_TEMPLATE_PATH", raising=False)
    output = tmp_path / "bid.zip"
    parts = [
        ExportPart(
            code="tech",
            title="技术标",
            chapters=[ExportChapter(title="实施计划", plain_text="项目实施内容。")],
            attachments=[
                ExportAttachment(
                    filename="技术附件.txt",
                    content_type="text/plain",
                    content_base64=_b64("技术附件内容"),
                )
            ],
        ),
        ExportPart(
            code="business",
            title="商务标",
            chapters=[ExportChapter(title="报价说明", plain_text="商务报价内容。")],
        ),
    ]

    export_bid_zip(
        "智慧交通平台",
        parts,
        output,
        attachments=[
            ExportAttachment(
                filename="资质证明.txt",
                content_type="text/plain",
                content_base64=_b64("资质证明内容"),
            )
        ],
        boq_files=[
            ExportAttachment(
                filename="工程量清单.xlsx",
                content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                content_base64=_b64("boq"),
            )
        ],
    )

    with ZipFile(output) as archive:
        names = sorted(archive.namelist())
        assert names == [
            "01_投标文件/01-技术标.docx",
            "01_投标文件/02-商务标.docx",
            "02_附件/tech/技术附件.txt",
            "02_附件/资质证明.txt",
            "03_工程量清单/工程量清单.xlsx",
            "manifest.json",
        ]
        first_docx = archive.read("01_投标文件/01-技术标.docx")
        manifest = json.loads(archive.read("manifest.json").decode("utf-8"))

    with ZipFile(BytesIO(first_docx)) as document_archive:
        document_xml = document_archive.read("word/document.xml").decode("utf-8")
        settings_xml = document_archive.read("word/settings.xml").decode("utf-8")

    assert manifest["bid_title"] == "智慧交通平台"
    assert manifest["e_bidding_structure"] == "standard"
    assert manifest["part_count"] == 2
    assert manifest["parts"][0]["filename"] == "01_投标文件/01-技术标.docx"
    assert manifest["parts"][0]["chapter_count"] == 1
    assert len(manifest["parts"][0]["sha256"]) == 64
    assert manifest["parts"][0]["attachments"][0]["filename"] == "02_附件/tech/技术附件.txt"
    assert manifest["attachments"][0]["filename"] == "02_附件/资质证明.txt"
    assert manifest["boq_files"][0]["filename"] == "03_工程量清单/工程量清单.xlsx"
    assert "TOC" in document_xml
    assert "实施计划" in document_xml
    assert "updateFields" in settings_xml


def test_export_bid_docx_renders_docxtpl_jinja_loops(tmp_path, monkeypatch) -> None:
    template_path = tmp_path / "jinja-template.docx"
    template = Document()
    template.add_paragraph("章节清单：{% for chapter in chapters %}{{ chapter.title }};{% endfor %}")
    template.save(template_path)
    monkeypatch.setenv("BID_EXPORT_TEMPLATE_PATH", str(template_path))
    output = tmp_path / "templated-loop.docx"

    export_bid_docx(
        "智慧交通平台",
        "技术标",
        [
            ExportChapter(title="实施计划", plain_text="正文一。"),
            ExportChapter(title="质量保证", plain_text="正文二。"),
        ],
        output,
        layout=ExportLayoutOptions(include_cover=False, include_toc=False, render_body=False),
    )

    document = Document(output)
    paragraph_text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    assert "章节清单：实施计划;质量保证;" in paragraph_text
    assert "正文一。" not in paragraph_text


def test_export_bid_docx_renders_template_placeholders_and_body_anchor(tmp_path, monkeypatch) -> None:
    template_path = tmp_path / "template.docx"
    template = Document()
    template.add_paragraph("项目：{{ bid_title }}")
    template.add_paragraph("分册：{{ part_title }}")
    template.add_paragraph("日期：{{ generated_at }}")
    template.add_paragraph("{{ZBT_BODY}}")
    template.save(template_path)
    monkeypatch.setenv("BID_EXPORT_TEMPLATE_PATH", str(template_path))
    output = tmp_path / "templated.docx"

    export_bid_docx(
        "智慧交通平台",
        "技术标",
        [ExportChapter(title="实施计划", plain_text="模板正文内容。")],
        output,
        layout=ExportLayoutOptions(include_cover=False, include_toc=False, generated_at="2026-06-12"),
    )

    document = Document(output)
    paragraph_text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    assert "项目：智慧交通平台" in paragraph_text
    assert "分册：技术标" in paragraph_text
    assert "日期：2026-06-12" in paragraph_text
    assert "{{ZBT_BODY}}" not in paragraph_text
    assert "实施计划" in paragraph_text
    assert "模板正文内容。" in paragraph_text


def test_template_field_replacement_stringifies_non_string_context() -> None:
    document = Document()
    document.add_paragraph("轮次：{{ review_round }}")
    document.add_paragraph("密封：{{ sealed }}")
    document.add_paragraph("备注：{{ note }}")

    _replace_template_fields(document, {"review_round": 2, "sealed": True, "note": None})
    paragraph_text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    assert "轮次：2" in paragraph_text
    assert "密封：True" in paragraph_text
    assert "备注：" in paragraph_text


def test_export_bid_pdf_reuses_master_docx_before_conversion(tmp_path, monkeypatch) -> None:
    monkeypatch.delenv("BID_EXPORT_TEMPLATE_PATH", raising=False)
    fake_soffice = tmp_path / "fake-soffice.sh"
    fake_soffice.write_text(
        """#!/bin/sh
outdir=""
input=""
while [ "$#" -gt 0 ]; do
  if [ "$1" = "--outdir" ]; then
    shift
    outdir="$1"
  else
    input="$1"
  fi
  shift
done
cp "$input" "$outdir/source.pdf"
""",
        encoding="utf-8",
    )
    fake_soffice.chmod(0o755)
    monkeypatch.setenv("LIBREOFFICE_PATH", str(fake_soffice))
    output = tmp_path / "bid.pdf"

    export_bid_pdf(
        "智慧交通平台",
        "技术标",
        [ExportChapter(title="实施计划", plain_text="项目实施内容。")],
        output,
        layout=ExportLayoutOptions(validate_pdf=False),
    )

    assert output.exists()
    assert "TOC" in _docx_xml(output, "word/document.xml")


def test_validate_pdf_output_checks_text_and_rendered_pixels(tmp_path) -> None:
    valid_pdf = tmp_path / "valid.pdf"
    pdf = fitz.open()
    page = pdf.new_page()
    page.insert_text((72, 72), "PDF validation sample")
    pdf.save(valid_pdf)
    pdf.close()

    result = _validate_pdf_output(valid_pdf)

    assert result["page_count"] == 1
    assert result["sampled_text_length"] > 0


def test_validate_pdf_output_rejects_blank_pdf(tmp_path) -> None:
    blank_pdf = tmp_path / "blank.pdf"
    pdf = fitz.open()
    pdf.new_page()
    pdf.save(blank_pdf)
    pdf.close()

    with pytest.raises(RuntimeError, match="no extractable text"):
        _validate_pdf_output(blank_pdf)


def test_attachment_local_path_is_rejected() -> None:
    attachment = ExportAttachment(filename="secret.txt", local_path="/etc/passwd")

    with pytest.raises(RuntimeError, match="local_path is not allowed"):
        _attachment_content(attachment)


def test_zip_path_removes_parent_segments() -> None:
    assert _safe_zip_path("../../evil/../资质.txt") == "evil/资质.txt"


def _b64(value: str) -> str:
    return base64.b64encode(value.encode("utf-8")).decode("ascii")


def _package_xml(path) -> str:
    with ZipFile(path) as archive:
        return "\n".join(
            archive.read(name).decode("utf-8", errors="ignore")
            for name in archive.namelist()
            if name.endswith(".xml")
        )


def _docx_xml(path, name: str) -> str:
    with ZipFile(path) as archive:
        return archive.read(name).decode("utf-8")
