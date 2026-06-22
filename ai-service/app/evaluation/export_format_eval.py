from __future__ import annotations

import argparse
import copy
import hashlib
import json
from pathlib import Path
from tempfile import TemporaryDirectory
from typing import Any
from zipfile import BadZipFile, ZipFile

import fitz
from docx import Document

from app.pipelines.export.docx_exporter import export_bid_docx, export_bid_pdf, export_bid_zip
from app.schemas.export import ExportAttachment, ExportChapter, ExportLayoutOptions, ExportPart


def evaluate_export_format(spec_path: Path, *, require_pdf: bool = False) -> dict[str, Any]:
    spec = json.loads(spec_path.read_text(encoding="utf-8"))
    if require_pdf:
        spec = copy.deepcopy(spec)
        pdf_spec = _dict_value(spec.get("pdf"))
        pdf_spec["enabled"] = True
        pdf_spec["allow_skip"] = False
        spec["pdf"] = pdf_spec
    checks: list[dict[str, Any]] = []
    with TemporaryDirectory() as tmpdir:
        tmp_path = Path(tmpdir)
        docx_result = _evaluate_docx(spec, tmp_path, checks)
        zip_result = _evaluate_zip(spec, tmp_path, checks)
        pdf_result = _evaluate_pdf(spec, tmp_path, checks)
    passed = sum(1 for check in checks if check["passed"])
    total = len(checks)
    return {
        "name": spec.get("name") or spec_path.stem,
        "status": "passed" if total and passed == total else "failed",
        "passed_checks": passed,
        "failed_checks": total - passed,
        "total_checks": total,
        "docx": docx_result,
        "zip": zip_result,
        "pdf": pdf_result,
        "checks": checks,
    }


def _evaluate_docx(spec: dict[str, Any], tmp_path: Path, checks: list[dict[str, Any]]) -> dict[str, Any]:
    docx_spec = _dict_value(spec.get("docx"))
    output = tmp_path / "export.docx"
    export_bid_docx(
        _text(spec.get("bid_title"), "智标通投标文件"),
        _text(spec.get("part_title"), "技术标"),
        _chapters(spec),
        output,
        layout=_layout(spec),
    )
    package = _docx_package(output)
    document_text = _docx_text(output)
    header_footer_text = _docx_header_footer_text(output)
    layout = _layout(spec)
    bid_title = _text(spec.get("bid_title"), "智标通投标文件")
    part_title = _text(spec.get("part_title"), "技术标")
    result = {
        "path": str(output),
        "size_bytes": output.stat().st_size,
        "table_count": _docx_table_count(output),
    }
    _add_check(checks, "export.docx.openable", bool(package), "valid docx zip", output.name)
    _add_check(checks, "export.docx.non_empty", output.stat().st_size > 0, ">0 bytes", output.stat().st_size)
    for text in _list_value(docx_spec.get("required_text")):
        expected = str(text)
        _add_check(
            checks,
            f"export.docx.text.{_check_slug(expected)}",
            expected in document_text,
            expected,
            _short_actual(document_text),
        )
    min_tables = _int_value(docx_spec.get("min_tables"), 0)
    if min_tables > 0:
        _add_check(checks, "export.docx.tables", result["table_count"] >= min_tables, f">={min_tables}", result["table_count"])
    package_xml = "\n".join(package.values())
    settings_xml = package.get("word/settings.xml", "")
    if docx_spec.get("require_cover", True) and layout.include_cover:
        cover_values = [bid_title, part_title, "投标文件"]
        missing_cover = [value for value in cover_values if value not in document_text]
        _add_check(checks, "export.docx.cover", not missing_cover, cover_values, missing_cover or "present")
    watermark_text = layout.watermark_text
    if docx_spec.get("require_watermark", bool(watermark_text)):
        _add_check(
            checks,
            "export.docx.watermark",
            bool(watermark_text) and watermark_text in package_xml,
            watermark_text or "watermark text configured",
            watermark_text in package_xml if watermark_text else "missing watermark text",
        )
    if docx_spec.get("require_toc", True):
        _add_check(checks, "export.docx.toc_field", "TOC" in package_xml, "TOC field", "TOC" in package_xml)
    if docx_spec.get("require_update_fields", True):
        _add_check(
            checks,
            "export.docx.update_fields",
            "updateFields" in settings_xml,
            "updateFields in settings.xml",
            "updateFields" in settings_xml,
        )
    if docx_spec.get("require_page_fields", True):
        has_page_fields = "PAGE" in package_xml and "NUMPAGES" in package_xml
        _add_check(checks, "export.docx.page_fields", has_page_fields, "PAGE and NUMPAGES fields", has_page_fields)
    if docx_spec.get("require_header_footer", True):
        has_header = any(name.startswith("word/header") for name in package)
        has_footer = any(name.startswith("word/footer") for name in package)
        _add_check(checks, "export.docx.header_footer", has_header and has_footer, "header and footer parts", {"header": has_header, "footer": has_footer})
        expected_header = layout.header_text or f"{bid_title} - {part_title}"
        expected_footer = layout.footer_text or "智标通投标文件导出"
        missing_header_footer = [
            value for value in (expected_header, expected_footer) if value not in header_footer_text
        ]
        _add_check(
            checks,
            "export.docx.header_footer_text",
            not missing_header_footer,
            [expected_header, expected_footer],
            missing_header_footer or "present",
        )
    return result


def _evaluate_zip(spec: dict[str, Any], tmp_path: Path, checks: list[dict[str, Any]]) -> dict[str, Any]:
    zip_spec = _dict_value(spec.get("zip"))
    if zip_spec.get("enabled") is False:
        return {"status": "disabled"}
    output = tmp_path / "export.zip"
    parts = _parts(spec)
    export_bid_zip(
        _text(spec.get("bid_title"), "智标通投标文件"),
        parts,
        output,
        layout=_layout(spec),
        attachments=_attachments(zip_spec.get("attachments")),
        boq_files=_attachments(zip_spec.get("boq_files")),
    )
    with ZipFile(output) as archive:
        names = sorted(archive.namelist())
        entry_bytes = {name: archive.read(name) for name in names if not name.endswith("/")}
        manifest = json.loads(archive.read("manifest.json").decode("utf-8")) if "manifest.json" in names else {}
        docx_entries = [name for name in names if name.endswith(".docx")]
        docx_openable = all(_zip_docx_openable(archive, name) for name in docx_entries)
    manifest_issues = _manifest_integrity_issues(
        manifest,
        entry_bytes,
        parts,
        _attachments(zip_spec.get("attachments")),
        _attachments(zip_spec.get("boq_files")),
        _text(spec.get("bid_title"), "智标通投标文件"),
        _layout(spec),
    )
    unsafe_paths = _unsafe_zip_paths(names)
    result = {
        "path": str(output),
        "size_bytes": output.stat().st_size,
        "entry_count": len(names),
        "docx_entry_count": len(docx_entries),
        "manifest_issues": manifest_issues,
    }
    _add_check(checks, "export.zip.openable", output.stat().st_size > 0, ">0 bytes", output.stat().st_size)
    if zip_spec.get("require_manifest", True):
        _add_check(checks, "export.zip.manifest", bool(manifest), "manifest.json", "manifest.json" in names)
        _add_check(checks, "export.zip.manifest_part_count", manifest.get("part_count") == len(parts), len(parts), manifest.get("part_count"))
        _add_check(checks, "export.zip.manifest_integrity", not manifest_issues, "manifest matches archive bytes and spec", manifest_issues or "ok")
    _add_check(checks, "export.zip.safe_paths", not unsafe_paths, "no absolute or parent zip paths", unsafe_paths or "ok")
    min_docx_entries = _int_value(zip_spec.get("min_docx_entries"), len(parts))
    _add_check(checks, "export.zip.docx_entries", len(docx_entries) >= min_docx_entries, f">={min_docx_entries}", len(docx_entries))
    _add_check(checks, "export.zip.docx_entries_openable", docx_openable, "all docx entries open", docx_entries)
    for entry in _list_value(zip_spec.get("expected_entries")):
        expected = str(entry)
        _add_check(checks, f"export.zip.entry.{_check_slug(expected)}", expected in names, expected, names)
    return result


def _evaluate_pdf(spec: dict[str, Any], tmp_path: Path, checks: list[dict[str, Any]]) -> dict[str, Any]:
    pdf_spec = _dict_value(spec.get("pdf"))
    if pdf_spec.get("enabled") is False:
        return {"status": "disabled"}
    output = tmp_path / "export.pdf"
    allow_skip = bool(pdf_spec.get("allow_skip", False))
    try:
        export_bid_pdf(
            _text(spec.get("bid_title"), "智标通投标文件"),
            _text(spec.get("part_title"), "技术标"),
            _chapters(spec),
            output,
            layout=_layout(spec),
        )
    except RuntimeError as exc:
        skipped = allow_skip and "LibreOffice" in str(exc)
        _add_check(checks, "export.pdf.generated", skipped, "generated PDF or explicit skip", "skipped: LibreOffice unavailable" if skipped else str(exc))
        return {"status": "skipped" if skipped else "failed", "reason": str(exc)}
    with fitz.open(output) as document:
        page_count = document.page_count
        text = "\n".join(page.get_text("text") for page in document)
        first_page_nonblank = _first_page_nonblank(document)
    result = {
        "status": "generated",
        "path": str(output),
        "size_bytes": output.stat().st_size,
        "page_count": page_count,
        "text_chars": len(text.strip()),
        "first_page_nonblank": first_page_nonblank,
    }
    min_pages = _int_value(pdf_spec.get("min_pages"), 1)
    _add_check(checks, "export.pdf.generated", output.stat().st_size > 0, ">0 bytes", output.stat().st_size)
    _add_check(checks, "export.pdf.openable", page_count >= min_pages, f">={min_pages} page(s)", page_count)
    if pdf_spec.get("require_text", True):
        _add_check(checks, "export.pdf.text_layer", bool(text.strip()), "non-empty text layer", len(text.strip()))
    if pdf_spec.get("require_first_page_nonblank", True):
        _add_check(checks, "export.pdf.first_page_nonblank", first_page_nonblank, True, first_page_nonblank)
    for text_value in _list_value(pdf_spec.get("required_text") or _dict_value(spec.get("docx")).get("required_text")):
        expected = str(text_value)
        matched = _pdf_text_contains(text, expected)
        _add_check(
            checks,
            f"export.pdf.text.{_check_slug(expected)}",
            matched,
            expected,
            _short_actual(text),
        )
    return result


def _docx_package(path: Path) -> dict[str, str]:
    try:
        with ZipFile(path) as archive:
            return {
                name: archive.read(name).decode("utf-8", errors="ignore")
                for name in archive.namelist()
                if name.endswith(".xml")
            }
    except BadZipFile:
        return {}


def _docx_text(path: Path) -> str:
    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def _docx_header_footer_text(path: Path) -> str:
    document = Document(path)
    parts: list[str] = []
    for section in document.sections:
        containers = [
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
            section.even_page_header,
            section.even_page_footer,
        ]
        for container in containers:
            parts.extend(paragraph.text for paragraph in container.paragraphs)
            for table in container.tables:
                for row in table.rows:
                    parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def _docx_table_count(path: Path) -> int:
    return len(Document(path).tables)


def _zip_docx_openable(archive: ZipFile, name: str) -> bool:
    try:
        with ZipFile(_bytes_buffer(archive.read(name))) as docx_archive:
            return "word/document.xml" in docx_archive.namelist()
    except BadZipFile:
        return False


def _bytes_buffer(value: bytes):
    from io import BytesIO

    return BytesIO(value)


def _first_page_nonblank(document: fitz.Document) -> bool:
    if document.page_count <= 0:
        return False
    pixmap = document[0].get_pixmap(matrix=fitz.Matrix(0.5, 0.5), alpha=False)
    samples = pixmap.samples
    if not samples:
        return False
    step = max(len(samples) // 12000, 1)
    return any(byte < 245 for byte in samples[::step])


def _manifest_integrity_issues(
    manifest: dict[str, Any],
    entry_bytes: dict[str, bytes],
    parts: list[ExportPart],
    attachments: list[ExportAttachment],
    boq_files: list[ExportAttachment],
    bid_title: str,
    layout: ExportLayoutOptions,
) -> list[str]:
    if not isinstance(manifest, dict) or not manifest:
        return ["manifest.json missing"]
    issues: list[str] = []
    if manifest.get("bid_title") != bid_title:
        issues.append("manifest bid_title mismatch")
    if manifest.get("template_name") != layout.template_name:
        issues.append("manifest template_name mismatch")
    if manifest.get("e_bidding_structure") != layout.e_bidding_structure:
        issues.append("manifest e_bidding_structure mismatch")
    if layout.generated_at and manifest.get("generated_at") != layout.generated_at:
        issues.append("manifest generated_at mismatch")
    elif not manifest.get("generated_at"):
        issues.append("manifest generated_at missing")
    manifest_parts = _dict_items(manifest.get("parts"))
    if len(manifest_parts) != len(parts):
        issues.append(f"manifest parts count mismatch: {len(manifest_parts)} != {len(parts)}")
    for index, part in enumerate(parts):
        if index >= len(manifest_parts):
            break
        record = manifest_parts[index]
        if record.get("code") != part.code:
            issues.append(f"manifest part {index + 1} code mismatch")
        if record.get("title") != part.title:
            issues.append(f"manifest part {index + 1} title mismatch")
        if record.get("chapter_count") != len(part.chapters):
            issues.append(f"manifest part {part.code} chapter_count mismatch")
        issues.extend(_manifest_entry_integrity_issues(record, entry_bytes, f"part {part.code}"))
        part_attachments = _dict_items(record.get("attachments"))
        if len(part_attachments) != len(part.attachments):
            issues.append(f"manifest part {part.code} attachment count mismatch")
        issues.extend(_manifest_entry_list_integrity_issues(part_attachments, entry_bytes, f"part {part.code} attachment"))
    manifest_attachments = _dict_items(manifest.get("attachments"))
    if len(manifest_attachments) != len(attachments):
        issues.append("manifest top-level attachment count mismatch")
    issues.extend(_manifest_entry_list_integrity_issues(manifest_attachments, entry_bytes, "attachment"))
    manifest_boq_files = _dict_items(manifest.get("boq_files"))
    if len(manifest_boq_files) != len(boq_files):
        issues.append("manifest BOQ file count mismatch")
    issues.extend(_manifest_entry_list_integrity_issues(manifest_boq_files, entry_bytes, "boq"))
    return issues


def _manifest_entry_list_integrity_issues(
    records: list[dict[str, Any]],
    entry_bytes: dict[str, bytes],
    label: str,
) -> list[str]:
    issues: list[str] = []
    for index, record in enumerate(records, start=1):
        issues.extend(_manifest_entry_integrity_issues(record, entry_bytes, f"{label} {index}"))
    return issues


def _manifest_entry_integrity_issues(
    record: dict[str, Any],
    entry_bytes: dict[str, bytes],
    label: str,
) -> list[str]:
    filename = str(record.get("filename") or "")
    if not filename:
        return [f"manifest {label} filename missing"]
    content = entry_bytes.get(filename)
    if content is None:
        return [f"manifest {label} file missing from zip: {filename}"]
    issues: list[str] = []
    if record.get("size_bytes") != len(content):
        issues.append(f"manifest {label} size mismatch")
    expected_sha = str(record.get("sha256") or "")
    actual_sha = hashlib.sha256(content).hexdigest()
    if expected_sha != actual_sha:
        issues.append(f"manifest {label} sha256 mismatch")
    return issues


def _unsafe_zip_paths(names: list[str]) -> list[str]:
    unsafe: list[str] = []
    for name in names:
        normalized = name.replace("\\", "/")
        parts = [part for part in normalized.split("/") if part]
        if normalized.startswith("/") or any(part == ".." for part in parts):
            unsafe.append(name)
    return unsafe


def _layout(spec: dict[str, Any]) -> ExportLayoutOptions:
    return ExportLayoutOptions(**_dict_value(spec.get("layout")))


def _chapters(spec: dict[str, Any]) -> list[ExportChapter]:
    return [ExportChapter(**item) for item in _dict_items(spec.get("chapters"))]


def _parts(spec: dict[str, Any]) -> list[ExportPart]:
    parts = _dict_items(_dict_value(spec.get("zip")).get("parts"))
    if parts:
        return [ExportPart(**item) for item in parts]
    return [
        ExportPart(
            code="main",
            title=_text(spec.get("part_title"), "技术标"),
            chapters=_chapters(spec),
        )
    ]


def _attachments(value: Any) -> list[ExportAttachment]:
    return [ExportAttachment(**item) for item in _dict_items(value)]


def _dict_value(value: Any) -> dict[str, Any]:
    return value if isinstance(value, dict) else {}


def _dict_items(value: Any) -> list[dict[str, Any]]:
    if not isinstance(value, list):
        return []
    return [item for item in value if isinstance(item, dict)]


def _list_value(value: Any) -> list[Any]:
    return value if isinstance(value, list) else []


def _int_value(value: Any, default: int) -> int:
    try:
        return int(value)
    except (TypeError, ValueError):
        return default


def _text(value: Any, default: str = "") -> str:
    text = "" if value is None else str(value).strip()
    return text or default


def _check_slug(value: str) -> str:
    slug = "".join(char if char.isalnum() else "-" for char in value.strip().lower())
    return slug.strip("-")[:48] or "value"


def _short_actual(value: str) -> str:
    normalized = " ".join(value.split())
    return normalized[:220]


def _pdf_text_contains(text: str, expected: str) -> bool:
    if expected in text:
        return True
    normalized_text = "".join(text.split())
    normalized_expected = "".join(expected.split())
    if not normalized_expected:
        return True
    if normalized_expected in normalized_text:
        return True
    significant = [char for char in normalized_expected if not char.isascii() or char.isalnum()]
    if len(significant) < 4:
        return False
    matched = sum(1 for char in significant if char in normalized_text)
    return matched / len(significant) >= 0.8


def _add_check(checks: list[dict[str, Any]], name: str, passed: bool, expected: Any, actual: Any) -> None:
    checks.append({"name": name, "passed": bool(passed), "expected": expected, "actual": actual})


def main() -> int:
    parser = argparse.ArgumentParser(description="Evaluate DOCX/PDF/ZIP export format regression.")
    parser.add_argument("--input", required=True, type=Path, help="Path to export format JSON.")
    parser.add_argument("--json", action="store_true", help="Print full JSON result.")
    parser.add_argument("--require-pdf", action="store_true", help="Fail instead of skipping PDF conversion.")
    args = parser.parse_args()

    result = evaluate_export_format(args.input, require_pdf=args.require_pdf)
    if args.json:
        print(json.dumps(result, ensure_ascii=False, indent=2))
    else:
        print(f"{result['status']} passed={result['passed_checks']}/{result['total_checks']} name={result['name']}")
        for check in result["checks"]:
            if not check["passed"]:
                print(f"- {check['name']}: expected={check['expected']!r} actual={check['actual']!r}")
    return 0 if result["status"] == "passed" else 1


if __name__ == "__main__":
    raise SystemExit(main())
