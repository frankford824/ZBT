from __future__ import annotations

import base64
import binascii
import hashlib
import json
import os
import re
import shutil
import subprocess
from io import BytesIO
from copy import deepcopy
from datetime import datetime, timezone
from pathlib import Path
from tempfile import TemporaryDirectory
from zipfile import ZIP_DEFLATED, ZipFile
from xml.sax.saxutils import escape

import fitz
from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.section import Section
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph
from docxtpl import DocxTemplate

from app.schemas.export import ExportAttachment, ExportChapter, ExportLayoutOptions, ExportPart

MARKDOWN_TABLE_SEPARATOR = re.compile(r"^:?-{3,}:?$")
TEMPLATE_FIELD = re.compile(r"{{\s*([A-Za-z0-9_.-]+)\s*}}")
CONTROL_CHARS = re.compile(r"[\x00-\x1f\x7f]")
WINDOWS_DRIVE_SEGMENT = re.compile(r"^[A-Za-z]:$")
WINDOWS_RESERVED_FILENAMES = {
    "CON",
    "PRN",
    "AUX",
    "NUL",
    *(f"COM{index}" for index in range(1, 10)),
    *(f"LPT{index}" for index in range(1, 10)),
}
ANCHOR_ALIASES = {
    "cover": {"ZBT_COVER", "zbt_cover", "cover"},
    "toc": {"ZBT_TOC", "zbt_toc", "toc"},
    "body": {"ZBT_BODY", "zbt_body", "body"},
}


def export_minimal_docx(title: str, paragraphs: list[str], output_path: Path) -> Path:
    document = Document()
    document.add_heading(title, level=1)
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    document.save(output_path)
    return output_path


def export_bid_docx(
    title: str,
    part_title: str,
    chapters: list[ExportChapter],
    output_path: Path,
    layout: ExportLayoutOptions | None = None,
) -> Path:
    layout = layout or ExportLayoutOptions()
    context = _template_context(title, part_title, layout, chapters=chapters)
    document = _new_document(context)
    _apply_master_layout(document, title, part_title, layout)
    anchors = _replace_template_fields(document, context)
    if layout.include_cover:
        _render_cover(document, title, part_title, layout, anchor=anchors.get("cover"))
    else:
        _remove_anchor(anchors.get("cover"))
    if layout.include_toc:
        _render_toc(document, layout, anchor=anchors.get("toc"))
    else:
        _remove_anchor(anchors.get("toc"))
    if layout.render_body:
        _render_bid_body(document, part_title, chapters, anchor=anchors.get("body"))
    else:
        _remove_anchor(anchors.get("body"))
    _enable_field_updates(document)
    document.save(output_path)
    return output_path


def export_bid_zip(
    title: str,
    parts: list[ExportPart],
    output_path: Path,
    layout: ExportLayoutOptions | None = None,
    attachments: list[ExportAttachment] | None = None,
    boq_files: list[ExportAttachment] | None = None,
) -> Path:
    layout = layout or ExportLayoutOptions()
    attachments = attachments or []
    boq_files = boq_files or []
    with TemporaryDirectory() as tmpdir:
        tmp_path = Path(tmpdir)
        with ZipFile(output_path, mode="w", compression=ZIP_DEFLATED) as archive:
            manifest_parts: list[dict[str, object]] = []
            used_paths: set[str] = set()
            for index, part in enumerate(parts, start=1):
                docx_name = f"{index:02d}-{_safe_filename(part.title or part.code)}.docx"
                docx_path = tmp_path / docx_name
                export_bid_docx(title, part.title, part.chapters, docx_path, layout=layout)
                zip_name = _dedupe_zip_path(_zip_part_path(layout, docx_name), used_paths)
                archive.write(docx_path, zip_name)
                content = docx_path.read_bytes()
                part_attachments = _write_attachments(
                    archive,
                    part.attachments,
                    layout,
                    "attachment",
                    used_paths,
                    part_code=part.code,
                )
                manifest_parts.append(
                    {
                        "code": part.code,
                        "title": part.title,
                        "filename": zip_name,
                        "chapter_count": len(part.chapters),
                        "size_bytes": len(content),
                        "sha256": hashlib.sha256(content).hexdigest(),
                        "attachments": part_attachments,
                    }
                )
            manifest_attachments = _write_attachments(
                archive,
                attachments,
                layout,
                "attachment",
                used_paths,
            )
            manifest_boq_files = _write_attachments(
                archive,
                boq_files,
                layout,
                "boq",
                used_paths,
            )
            if layout.include_manifest:
                archive.writestr(
                    "manifest.json",
                    json.dumps(
                        {
                            "bid_title": title,
                            "template_name": layout.template_name,
                            "e_bidding_structure": layout.e_bidding_structure,
                            "generated_at": _generated_at(layout),
                            "part_count": len(manifest_parts),
                            "parts": manifest_parts,
                            "attachments": manifest_attachments,
                            "boq_files": manifest_boq_files,
                        },
                        ensure_ascii=False,
                        indent=2,
                    ),
                )
    return output_path


def export_bid_pdf(
    title: str,
    part_title: str,
    chapters: list[ExportChapter],
    output_path: Path,
    layout: ExportLayoutOptions | None = None,
) -> Path:
    layout = layout or ExportLayoutOptions()
    with TemporaryDirectory() as tmpdir:
        tmp_path = Path(tmpdir)
        docx_path = tmp_path / "source.docx"
        export_bid_docx(title, part_title, chapters, docx_path, layout=layout)
        soffice = _libreoffice_executable()
        try:
            completed = subprocess.run(
                [
                    soffice,
                    "--headless",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    str(tmp_path),
                    str(docx_path),
                ],
                check=False,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True,
                timeout=90,
            )
        except (OSError, subprocess.TimeoutExpired):
            raise RuntimeError("LibreOffice PDF conversion failed") from None
        pdf_path = docx_path.with_suffix(".pdf")
        if completed.returncode != 0 or not pdf_path.exists():
            raise RuntimeError("LibreOffice PDF conversion failed")
        if layout.validate_pdf:
            _validate_pdf_output(pdf_path)
        shutil.move(str(pdf_path), output_path)
    return output_path


def _libreoffice_executable() -> str:
    configured = os.getenv("LIBREOFFICE_PATH", "").strip()
    if configured:
        return configured
    executable = shutil.which("soffice") or shutil.which("libreoffice")
    if not executable:
        raise RuntimeError("LibreOffice executable not found; set LIBREOFFICE_PATH")
    return executable


def _new_document(context: dict[str, object]) -> DocxDocument:
    template_path = os.getenv("BID_EXPORT_TEMPLATE_PATH", "").strip()
    if not template_path:
        return Document()
    path = Path(template_path)
    if not path.exists():
        raise RuntimeError(f"BID_EXPORT_TEMPLATE_PATH does not exist: {template_path}")
    template = DocxTemplate(str(path))
    context_with_anchors = {
        "ZBT_COVER": "{{ZBT_COVER}}",
        "ZBT_TOC": "{{ZBT_TOC}}",
        "ZBT_BODY": "{{ZBT_BODY}}",
        **context,
    }
    template.render(context_with_anchors)
    rendered = BytesIO()
    template.save(rendered)
    rendered.seek(0)
    return Document(rendered)


def _template_context(
    title: str,
    part_title: str,
    layout: ExportLayoutOptions,
    chapters: list[ExportChapter],
) -> dict[str, object]:
    context = {
        "bid_title": title,
        "title": title,
        "part_title": part_title,
        "template_name": layout.template_name,
        "generated_at": _generated_at(layout),
        "document_type": "投标文件",
        "chapters": [
            {"title": chapter.title, "plain_text": chapter.plain_text}
            for chapter in chapters
        ],
    }
    context.update(_public_template_context(layout.context))
    return context


def _public_template_context(context: dict[str, object]) -> dict[str, object]:
    reserved = {alias for aliases in ANCHOR_ALIASES.values() for alias in aliases}
    return {key: value for key, value in context.items() if key not in reserved}


def _generated_at(layout: ExportLayoutOptions) -> str:
    return layout.generated_at or datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")


def _write_attachments(
    archive: ZipFile,
    attachments: list[ExportAttachment],
    layout: ExportLayoutOptions,
    category: str,
    used_paths: set[str],
    part_code: str | None = None,
) -> list[dict[str, object]]:
    manifest_items: list[dict[str, object]] = []
    for attachment in attachments:
        content = _attachment_content(attachment)
        zip_path = _dedupe_zip_path(_zip_attachment_path(layout, attachment, category, part_code), used_paths)
        archive.writestr(zip_path, content)
        manifest_items.append(
            {
                "filename": zip_path,
                "original_filename": attachment.filename,
                "category": attachment.category or category,
                "content_type": attachment.content_type,
                "object_key": attachment.object_key,
                "size_bytes": len(content),
                "sha256": hashlib.sha256(content).hexdigest(),
            }
        )
    return manifest_items


def _attachment_content(attachment: ExportAttachment) -> bytes:
    if attachment.content_base64:
        try:
            return base64.b64decode(attachment.content_base64, validate=True)
        except binascii.Error as exc:
            raise RuntimeError(f"attachment content_base64 is invalid: {attachment.filename}") from exc
    if attachment.local_path:
        raise RuntimeError("attachment local_path is not allowed")
    raise RuntimeError(f"attachment content missing: {attachment.filename}")


def _zip_part_path(layout: ExportLayoutOptions, filename: str) -> str:
    if layout.e_bidding_structure == "flat":
        return filename
    return f"01_投标文件/{filename}"


def _zip_attachment_path(
    layout: ExportLayoutOptions,
    attachment: ExportAttachment,
    category: str,
    part_code: str | None,
) -> str:
    filename = _safe_filename(attachment.filename)
    if attachment.zip_path:
        return _safe_zip_path(attachment.zip_path)
    if layout.e_bidding_structure == "flat":
        folder = "boq" if category == "boq" else "attachments"
        return f"{folder}/{filename}"
    if category == "boq":
        return f"03_工程量清单/{filename}"
    if part_code:
        return f"02_附件/{_safe_filename(part_code)}/{filename}"
    return f"02_附件/{filename}"


def _dedupe_zip_path(path: str, used_paths: set[str]) -> str:
    cleaned = _safe_zip_path(path)
    candidate = cleaned
    suffix = 2
    while candidate in used_paths:
        path_obj = Path(cleaned)
        stem = path_obj.stem
        extension = path_obj.suffix
        candidate = str(path_obj.with_name(f"{stem}-{suffix}{extension}")).replace("\\", "/")
        suffix += 1
    used_paths.add(candidate)
    return candidate


def _safe_zip_path(value: str) -> str:
    cleaned = value.strip().replace("\\", "/").strip("/")
    parts = []
    for part in cleaned.split("/"):
        raw_part = part.strip()
        if not raw_part or raw_part in {".", ".."} or WINDOWS_DRIVE_SEGMENT.fullmatch(raw_part):
            continue
        safe = _safe_filename(raw_part).strip(".")
        if safe:
            parts.append(safe)
    return "/".join(parts) or "attachment"


def _safe_filename(value: str) -> str:
    cleaned = CONTROL_CHARS.sub("-", value.strip())
    for char in '/\\:*?"<>|':
        cleaned = cleaned.replace(char, "-")
    cleaned = cleaned.strip(" .") or "document"
    return _avoid_windows_reserved_filename(cleaned)


def _avoid_windows_reserved_filename(value: str) -> str:
    stem = value.split(".", 1)[0].rstrip(" .").upper()
    if stem in WINDOWS_RESERVED_FILENAMES:
        return f"_{value}"
    return value


def _validate_pdf_output(path: Path) -> dict[str, object]:
    try:
        doc = fitz.open(path)
    except Exception as exc:
        raise RuntimeError(f"PDF validation failed: cannot open PDF: {exc}") from exc
    try:
        if doc.page_count <= 0:
            raise RuntimeError("PDF validation failed: no pages")
        text = "\n".join(
            doc[index].get_text("text").strip() for index in range(min(doc.page_count, 3))
        ).strip()
        if not text:
            raise RuntimeError("PDF validation failed: no extractable text")
        first_page = doc[0]
        pixmap = first_page.get_pixmap(matrix=fitz.Matrix(0.25, 0.25), alpha=False)
        samples = pixmap.samples
        stride = max(1, len(samples) // 4096)
        if len(set(samples[::stride])) < 2:
            raise RuntimeError("PDF validation failed: first page rendered blank")
        return {"page_count": doc.page_count, "sampled_text_length": len(text)}
    finally:
        doc.close()


def _replace_template_fields(
    document: DocxDocument,
    context: dict[str, object],
) -> dict[str, Paragraph]:
    anchors: dict[str, Paragraph] = {}
    for paragraph in _iter_paragraphs(document):
        text = paragraph.text
        if not text:
            continue
        for section, aliases in ANCHOR_ALIASES.items():
            if any(f"{{{{{alias}}}}}" in text or f"{{{{ {alias} }}}}" in text for alias in aliases):
                anchors[section] = paragraph
        replaced = TEMPLATE_FIELD.sub(lambda match: _template_replacement(context, match), text)
        if replaced != text:
            paragraph.text = replaced
    return anchors


def _template_replacement(context: dict[str, object], match: re.Match[str]) -> str:
    key = match.group(1)
    if key not in context:
        return match.group(0)
    value = context[key]
    if value is None:
        return ""
    return str(value)


def _iter_paragraphs(document: DocxDocument):
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
    for section in document.sections:
        containers = [
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
            section.even_page_header,
            section.even_page_footer,
        ]
        for container in containers:
            yield from container.paragraphs
            for table in container.tables:
                for row in table.rows:
                    for cell in row.cells:
                        yield from cell.paragraphs


def _insert_document_after(anchor: Paragraph, source: DocxDocument) -> None:
    body_elements = [
        deepcopy(element)
        for element in source.element.body
        if element.tag != qn("w:sectPr")
    ]
    for element in reversed(body_elements):
        anchor._p.addnext(element)
    parent = anchor._element.getparent()
    if parent is not None:
        parent.remove(anchor._element)


def _remove_anchor(anchor: Paragraph | None) -> None:
    if anchor is None:
        return
    parent = anchor._element.getparent()
    if parent is not None:
        parent.remove(anchor._element)


def _section_document() -> DocxDocument:
    document = Document()
    _apply_styles(document)
    return document


def _apply_master_layout(
    document: DocxDocument,
    title: str,
    part_title: str,
    layout: ExportLayoutOptions,
) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.6)
    section.bottom_margin = Cm(2.4)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1.2)
    section.different_first_page_header_footer = layout.include_cover
    _apply_styles(document)
    _render_header_footer(section, title, part_title, layout)


def _apply_styles(document: DocxDocument) -> None:
    _set_style_font(document, "Normal", "Times New Roman", "SimSun", 12, RGBColor(0x1F, 0x29, 0x37))
    normal = document.styles["Normal"]
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.first_line_indent = Cm(0.74)

    _set_style_font(document, "Title", "Times New Roman", "Microsoft YaHei", 22, RGBColor(0x11, 0x1B, 0x2E))
    _set_style_font(document, "Heading 1", "Times New Roman", "Microsoft YaHei", 18, RGBColor(0x11, 0x1B, 0x2E))
    _set_style_font(document, "Heading 2", "Times New Roman", "Microsoft YaHei", 15, RGBColor(0x1D, 0x4E, 0x89))
    _set_style_font(document, "Heading 3", "Times New Roman", "Microsoft YaHei", 13, RGBColor(0x1F, 0x29, 0x37))
    for name in ("Heading 1", "Heading 2", "Heading 3"):
        style = document.styles[name]
        style.paragraph_format.first_line_indent = None
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def _set_style_font(
    document: DocxDocument,
    style_name: str,
    ascii_font: str,
    east_asia_font: str,
    size: int,
    color: RGBColor,
) -> None:
    style = document.styles[style_name]
    style.font.name = ascii_font
    style.font.size = Pt(size)
    style.font.color.rgb = color
    if style.element.rPr is None:
        style.element.get_or_add_rPr()
    style.element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_font)


def _render_header_footer(
    section: Section,
    title: str,
    part_title: str,
    layout: ExportLayoutOptions,
) -> None:
    header_text = layout.header_text or f"{title} - {part_title}"
    header = section.header
    header.paragraphs[0].text = header_text
    header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _format_runs(header.paragraphs[0], font_size=9, color=RGBColor(0x55, 0x65, 0x74))

    watermark_text = layout.watermark_text or os.getenv("BID_EXPORT_WATERMARK_TEXT", "").strip()
    if watermark_text:
        _add_watermark(header, watermark_text)
        if section.different_first_page_header_footer:
            _add_watermark(section.first_page_header, watermark_text)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.text = layout.footer_text or "智标通投标文件导出"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if layout.include_page_numbers:
        paragraph.add_run("  |  第 ")
        _add_field(paragraph, "PAGE", "1")
        paragraph.add_run(" 页 / 共 ")
        _add_field(paragraph, "NUMPAGES", "1")
        paragraph.add_run(" 页")
    _format_runs(paragraph, font_size=9, color=RGBColor(0x55, 0x65, 0x74))


def _add_watermark(header, text: str) -> None:
    paragraph = header.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    watermark_xml = f"""
    <w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
         xmlns:v="urn:schemas-microsoft-com:vml"
         xmlns:o="urn:schemas-microsoft-com:office:office">
      <w:pict>
        <v:shape id="ZBTWatermark" o:spid="_x0000_s1025" type="#_x0000_t136"
          style="position:absolute;margin-left:0;margin-top:0;width:420pt;height:120pt;rotation:315;z-index:-251654144;mso-position-horizontal:center;mso-position-horizontal-relative:margin;mso-position-vertical:center;mso-position-vertical-relative:margin"
          fillcolor="#BFBFBF" stroked="f">
          <v:fill opacity=".14"/>
          <v:textpath style="font-family:&quot;Microsoft YaHei&quot;;font-size:1pt" string="{escape(text)}"/>
        </v:shape>
      </w:pict>
    </w:r>
    """
    paragraph._p.append(parse_xml(watermark_xml))


def _render_cover(
    document: DocxDocument,
    title: str,
    part_title: str,
    layout: ExportLayoutOptions,
    anchor: Paragraph | None = None,
) -> None:
    if anchor is not None:
        section = _section_document()
        _render_cover(section, title, part_title, layout)
        _insert_document_after(anchor, section)
        return
    if not layout.include_cover:
        document.add_heading(title, level=1)
        document.add_heading(part_title, level=2)
        return
    for _ in range(4):
        document.add_paragraph()
    title_paragraph = document.add_paragraph(style="Title")
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.add_run(title).bold = True
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(part_title)
    subtitle_run.bold = True
    subtitle_run.font.size = Pt(18)
    document_type = document.add_paragraph()
    document_type.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document_type.add_run("投标文件").font.size = Pt(16)
    if layout.generated_at:
        generated = document.add_paragraph()
        generated.alignment = WD_ALIGN_PARAGRAPH.CENTER
        generated.add_run(layout.generated_at).font.size = Pt(11)
    document.add_page_break()


def _render_toc(
    document: DocxDocument,
    layout: ExportLayoutOptions,
    anchor: Paragraph | None = None,
) -> None:
    if anchor is not None:
        section = _section_document()
        _render_toc(section, layout)
        _insert_document_after(anchor, section)
        return
    if not layout.include_toc:
        return
    heading = document.add_paragraph("目录", style="Heading 1")
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = None
    _add_field(paragraph, r'TOC \o "1-3" \h \z \u', "目录")
    document.add_page_break()


def _render_bid_body(
    document: DocxDocument,
    part_title: str,
    chapters: list[ExportChapter],
    anchor: Paragraph | None = None,
) -> None:
    if anchor is not None:
        section = _section_document()
        _render_bid_body(section, part_title, chapters)
        _insert_document_after(anchor, section)
        return
    document.add_heading(part_title, level=1)
    for chapter in chapters:
        document.add_heading(chapter.title, level=2)
        text = chapter.plain_text.strip() or "本章节暂无内容，需要人工补充。"
        _render_plain_text(document, text)


def _render_plain_text(document: DocxDocument, text: str) -> None:
    lines = text.splitlines()
    index = 0
    while index < len(lines):
        stripped = lines[index].strip()
        if not stripped:
            index += 1
            continue
        table_rows, next_index = _consume_markdown_table(lines, index)
        if table_rows:
            _add_table(document, table_rows)
            index = next_index
            continue
        heading_level, heading_text = _markdown_heading(stripped)
        if heading_level:
            document.add_heading(heading_text, level=min(3, heading_level + 2))
        elif ordered := re.match(r"^\d+[.)、]\s+(.+)$", stripped):
            document.add_paragraph(ordered.group(1), style=_style_or_normal(document, "List Number"))
        elif bullet := re.match(r"^[-*•]\s+(.+)$", stripped):
            document.add_paragraph(bullet.group(1), style=_style_or_normal(document, "List Bullet"))
        else:
            document.add_paragraph(stripped)
        index += 1


def _markdown_heading(line: str) -> tuple[int | None, str]:
    match = re.match(r"^(#{1,3})\s+(.+)$", line)
    if not match:
        return None, line
    return len(match.group(1)), match.group(2).strip()


def _consume_markdown_table(lines: list[str], start: int) -> tuple[list[list[str]] | None, int]:
    if "|" not in lines[start]:
        return None, start
    raw_rows: list[list[str]] = []
    index = start
    while index < len(lines) and "|" in lines[index]:
        cells = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
        if len(cells) < 2:
            break
        raw_rows.append(cells)
        index += 1
    rows = [row for row in raw_rows if not all(MARKDOWN_TABLE_SEPARATOR.match(cell) for cell in row)]
    width = max((len(row) for row in rows), default=0)
    if len(rows) < 2 or width < 2:
        return None, start
    return [row + [""] * (width - len(row)) for row in rows], index


def _add_table(document: DocxDocument, rows: list[list[str]]) -> None:
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = _table_style(document)
    for row_index, row in enumerate(rows):
        for column_index, cell_value in enumerate(row):
            cell = table.cell(row_index, column_index)
            cell.text = cell_value
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.first_line_indent = None
                if row_index == 0 and paragraph.runs:
                    paragraph.runs[0].bold = True
    document.add_paragraph()


def _table_style(document: DocxDocument) -> str:
    return "Table Grid" if "Table Grid" in document.styles else "Normal Table"


def _style_or_normal(document: DocxDocument, style_name: str) -> str:
    return style_name if style_name in document.styles else "Normal"


def _format_runs(paragraph, font_size: int, color: RGBColor) -> None:
    for run in paragraph.runs:
        run.font.size = Pt(font_size)
        run.font.color.rgb = color


def _add_field(paragraph, instruction: str, placeholder: str) -> None:
    begin = paragraph.add_run()
    begin_char = OxmlElement("w:fldChar")
    begin_char.set(qn("w:fldCharType"), "begin")
    begin._r.append(begin_char)

    instr = paragraph.add_run()
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction
    instr._r.append(instr_text)

    separate = paragraph.add_run()
    separate_char = OxmlElement("w:fldChar")
    separate_char.set(qn("w:fldCharType"), "separate")
    separate._r.append(separate_char)
    paragraph.add_run(placeholder)

    end = paragraph.add_run()
    end_char = OxmlElement("w:fldChar")
    end_char.set(qn("w:fldCharType"), "end")
    end._r.append(end_char)


def _enable_field_updates(document: DocxDocument) -> None:
    settings = document.settings.element
    for existing in settings.findall(qn("w:updateFields")):
        settings.remove(existing)
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)
